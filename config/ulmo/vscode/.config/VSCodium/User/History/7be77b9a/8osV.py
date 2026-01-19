from docx import Document

# Create a new Word document
doc = Document()
doc.add_heading("Вычисления выражения (-15*a + b - c/4) / (b*a - 8)", level=1)

doc.add_paragraph("Рассматриваем выражение:")
doc.add_paragraph("expr = (-15a + b - c/4) / (ba - 8)")

cases = [
    {
        "a": -4, "b": -5, "c": -20,
        "numerator": "(-15)*(-4) + (-5) - (-20)/4 = 60 - 5 + 5 = 60",
        "denominator": "(-5)*(-4) - 8 = 20 - 8 = 12",
        "result": "60 / 12 = 5.0",
        "res": "(60.0, 12)",
        "note": "Промежуточные результаты совпадают: числитель = 60.0, знаменатель = 12"
    },
    {
        "a": 38, "b": 4, "c": 40,
        "numerator": "(-15)*38 + 4 - 40/4 = -570 + 4 - 10 = -576",
        "denominator": "4*38 - 8 = 152 - 8 = 144",
        "result": "-576 / 144 = -4.0",
        "res": "(-576.0, 144)",
        "note": "Совпадает: (-576.0, 144)"
    },
    {
        "a": -2, "b": -3, "c": -12,
        "numerator": "(-15)*(-2) + (-3) - (-12)/4 = 30 - 3 + 3 = 30",
        "denominator": "(-3)*(-2) - 8 = 6 - 8 = -2",
        "result": "30 / -2 = -15.0",
        "res": "(30.0, -2)",
        "note": "Совпадает по частям: числитель 30.0, знаменатель -2"
    },
    {
        "a": 38, "b": -15, "c": -28,
        "numerator": "(-15)*38 + (-15) - (-28)/4 = -570 -15 + 7 = -578",
        "denominator": "-15*38 - 8 = -570 - 8 = -578",
        "result": "-578 / -578 = 1.0",
        "res": "(-578.0, -578)",
        "note": "Числитель и знаменатель равны -578"
    }
]

for i, case in enumerate(cases, 1):
    doc.add_heading(f"Случай {i}) a = {case['a']}, b = {case['b']}, c = {case['c']}", level=2)
    doc.add_paragraph(f"Числитель: {case['numerator']}")
    doc.add_paragraph(f"Знаменатель: {case['denominator']}")
    doc.add_paragraph(f"Результат: {case['result']}")
    doc.add_paragraph(f"res: {case['res']}")
    doc.add_paragraph(f"Примечание: {case['note']}")

doc.add_paragraph("\nВывод: Результаты res=(числитель, знаменатель) указаны именно как значения числителя и знаменателя, а не итог деления. Все вычисления полностью подтверждены.")

# Save the document
file_path = "word_res.docs"
doc.save(file_path)